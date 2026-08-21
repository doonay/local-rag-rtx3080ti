import pytest
from pathlib import Path

@pytest.fixture
def sample_pdf_path():
    """Создает тестовый PDF файл"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    pdf_path = Path("/tmp/test.pdf")
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "Test PDF Content")
    c.drawString(100, 730, "This is a test document for RAG system.")
    c.save()
    
    yield pdf_path
    
    # Cleanup
    if pdf_path.exists():
        pdf_path.unlink()

@pytest.fixture
def sample_docx_path():
    """Создает тестовый DOCX файл"""
    from docx import Document
    
    docx_path = Path("/tmp/test.docx")
    doc = Document()
    doc.add_paragraph("Test DOCX Content")
    doc.add_paragraph("This is a test document for RAG system.")
    doc.save(str(docx_path))
    
    yield docx_path
    
    if docx_path.exists():
        docx_path.unlink()